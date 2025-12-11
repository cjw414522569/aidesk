import os
import sys
import markdown
from docx import Document
from openpyxl import Workbook
from PyPDF2 import PdfReader
import config

def get_resource_path(relative_path):
    """获取资源文件的绝对路径，兼容开发环境和打包后的exe"""
    if getattr(sys, 'frozen', False):
        # 打包后的exe，资源在_MEIPASS目录
        base_path = sys._MEIPASS
    else:
        # 开发环境，资源在当前目录
        base_path = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_path, relative_path)

class FileHandler:
    def read_file(self, filepath):
        """读取文件内容，支持多种格式"""
        try:
            # 确保使用绝对路径
            filepath = os.path.abspath(filepath)
            
            # 检查文件是否存在
            if not os.path.exists(filepath):
                return f"文件不存在: {filepath}"
            
            ext = os.path.splitext(filepath)[1].lower()
            
            if ext == '.pdf':
                try:
                    reader = PdfReader(filepath)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text()
                    return text if text else "PDF文件内容为空或无法读取"
                except Exception as e:
                    return f"读取PDF文件失败: {str(e)}"
            
            elif ext in ['.txt', '.md']:
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        return f.read()
                except UnicodeDecodeError:
                    # 尝试其他编码
                    try:
                        with open(filepath, 'r', encoding='gbk') as f:
                            return f.read()
                    except UnicodeDecodeError:
                        return "文件编码不支持，请使用UTF-8或GBK编码"
                except Exception as e:
                    return f"读取文本文件失败: {str(e)}"
            
            elif ext == '.docx':
                try:
                    doc = Document(filepath)
                    return '\n'.join([p.text for p in doc.paragraphs])
                except Exception as e:
                    return f"读取Word文档失败: {str(e)}"
            
            else:
                return f"不支持的文件格式: {ext}，支持的格式有：.txt, .md, .docx, .pdf"
                
        except Exception as e:
            return f"读取文件失败: {str(e)}"
    
    def markdown_to_word(self, md_file, output_file):
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        html = markdown.markdown(md_content)
        doc = Document()
        doc.add_paragraph(html)
        doc.save(output_file)
        return f"已转换为：{output_file}"
    
    def markdown_to_excel(self, md_file, output_file):
        with open(md_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        wb = Workbook()
        ws = wb.active
        
        for i, line in enumerate(lines, 1):
            ws.cell(row=i, column=1, value=line.strip())
        
        wb.save(output_file)
        return f"已转换为：{output_file}"
    
    def create_folder(self, folder_path):
        os.makedirs(folder_path, exist_ok=True)
        return f"已创建文件夹：{folder_path}"
    
    def create_file(self, filepath, content=""):
        """创建文件并写入内容"""
        try:
            # 确保使用绝对路径
            filepath = os.path.abspath(filepath)
            
            # 创建目录（如果需要）
            dir_path = os.path.dirname(filepath)
            if dir_path:
                os.makedirs(dir_path, exist_ok=True)
            
            # 创建文件
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return f"已创建文件：{filepath}"
        except Exception as e:
            return f"创建文件失败: {str(e)}"